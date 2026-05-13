"""Word export for the completeness checklist."""

from __future__ import annotations

import io
from datetime import date
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.db.models import ChecklistItem
from app.services.checklist_generator import CATEGORY_LABELS

CATEGORY_ICONS = {
    "missing_doc": "📄",
    "warning_note": "⚠️",
    "mortgage": "🏦",
    "lender": "💼",
    "signing": "✍️",
    "corporate": "🏢",
    "other": "📌",
}


def _rtl_para(doc: Document, text: str = "", style: str | None = None) -> Any:
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.insert(0, bidi)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    pPr.append(jc)
    if text:
        run = p.add_run(text)
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement("w:rtl")
        rPr.append(rtl)
    return p


def _add_heading(doc: Document, text: str, level: int = 1) -> Any:
    p = doc.add_heading(level=level)
    p.clear()
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.insert(0, bidi)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    pPr.append(jc)
    run = p.add_run(text)
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rPr.append(rtl)
    return p


def _checkbox(completed: bool) -> str:
    return "☑" if completed else "☐"


def generate_checklist_docx(items: list[ChecklistItem], project_title: str) -> bytes:
    doc = Document()

    # RTL defaults
    styles_el = doc.styles.element
    doc_defaults = styles_el.find(qn("w:docDefaults"))
    if doc_defaults is not None:
        pPrDef = doc_defaults.find(qn("w:pPrDefault"))
        if pPrDef is None:
            pPrDef = OxmlElement("w:pPrDefault")
            doc_defaults.insert(0, pPrDef)
        pPr = pPrDef.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            pPrDef.append(pPr)
        if pPr.find(qn("w:bidi")) is None:
            pPr.insert(0, OxmlElement("w:bidi"))

    # Title
    title_p = doc.add_heading(level=0)
    title_p.clear()
    pPr = title_p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.insert(0, bidi)
    run = title_p.add_run(f"רשימת השלמות — {project_title}")
    rPr = run._r.get_or_add_rPr()
    rPr.append(OxmlElement("w:rtl"))

    date_str = date.today().strftime("%d/%m/%Y")
    _rtl_para(doc, f"תאריך: {date_str}")

    # Stats
    total = len(items)
    done = sum(1 for i in items if i.is_completed)
    _rtl_para(doc, f"סה״כ פריטים: {total} | הושלמו: {done} | ממתינים: {total - done}")
    doc.add_paragraph()

    # Group by category
    from collections import defaultdict
    grouped: dict[str, list[ChecklistItem]] = defaultdict(list)
    for item in items:
        grouped[item.category].append(item)

    category_order = ["signing", "warning_note", "mortgage", "missing_doc", "lender", "corporate", "other"]
    for cat in category_order:
        cat_items = grouped.get(cat, [])
        if not cat_items:
            continue

        label = CATEGORY_LABELS.get(cat, cat)
        icon = CATEGORY_ICONS.get(cat, "•")
        _add_heading(doc, f"{icon}  {label}", level=1)

        # Table
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        # RTL table
        tblPr = table._tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            table._tbl.insert(0, tblPr)
        bidi_v = OxmlElement("w:bidiVisual")
        tblPr.append(bidi_v)

        # Set column widths (status, title, description)
        tbl_grid = OxmlElement("w:tblGrid")
        for w in [700, 3200, 5100]:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(w))
            tbl_grid.append(gc)
        table._tbl.insert(1, tbl_grid)

        # Header row
        hdr = table.rows[0].cells
        for idx, txt in enumerate(["סטטוס", "פעולה נדרשת", "פירוט"]):
            hdr[idx].text = ""
            p = hdr[idx].paragraphs[0]
            pPr = p._p.get_or_add_pPr()
            pPr.insert(0, OxmlElement("w:bidi"))
            run = p.add_run(txt)
            run.bold = True
            rPr = run._r.get_or_add_rPr()
            rPr.append(OxmlElement("w:rtl"))

        for item in cat_items:
            row = table.add_row().cells
            # Status cell
            row[0].text = ""
            p0 = row[0].paragraphs[0]
            p0._p.get_or_add_pPr().insert(0, OxmlElement("w:bidi"))
            r0 = p0.add_run(_checkbox(item.is_completed))
            r0.font.size = Pt(14)
            if item.is_completed:
                r0.font.color.rgb = RGBColor(0x22, 0xC5, 0x5E)  # green

            # Title cell
            row[1].text = ""
            p1 = row[1].paragraphs[0]
            p1._p.get_or_add_pPr().insert(0, OxmlElement("w:bidi"))
            r1 = p1.add_run(item.title)
            r1._r.get_or_add_rPr().append(OxmlElement("w:rtl"))
            if item.is_completed:
                r1.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)  # muted

            # Description cell
            row[2].text = ""
            p2 = row[2].paragraphs[0]
            p2._p.get_or_add_pPr().insert(0, OxmlElement("w:bidi"))
            r2 = p2.add_run(item.description or "")
            r2._r.get_or_add_rPr().append(OxmlElement("w:rtl"))
            r2.font.size = Pt(9)
            r2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
