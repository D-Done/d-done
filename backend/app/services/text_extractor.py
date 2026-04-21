"""Extract plain-text content from office documents for Gemini inline injection.

Gemini's document API natively understands PDFs and images, but not Excel, Word,
CSV, HTML, or email files. This module downloads such files from GCS and converts
them to a text representation that can be sent as ``Part.from_text()`` alongside
the native-PDF parts, so the model still "sees" the content of those files.

Limits (to protect context window and latency):
- Files larger than ``MAX_EXTRACT_BYTES`` are skipped (logged as WARNING).
- Each extracted text is capped at ``MAX_TEXT_CHARS`` characters.
"""

from __future__ import annotations

import asyncio
import io
import logging
from typing import NamedTuple

logger = logging.getLogger(__name__)

# Files larger than this will not be downloaded/extracted.
MAX_EXTRACT_BYTES = 10 * 1024 * 1024  # 10 MiB

# Maximum characters kept per extracted file (~30k tokens for most models).
MAX_TEXT_CHARS = 120_000

# System / junk files that carry no useful content for the DD pipeline.
_SKIP_BASENAMES: frozenset[str] = frozenset({
    ".DS_Store", "Thumbs.db", "desktop.ini", ".gitkeep", ".gitignore",
    ".gitmodules", "thumbs.db",
})
_SKIP_EXTENSIONS: frozenset[str] = frozenset({
    "ds_store", "gitignore", "gitkeep", "lnk", "tmp", "bak",
})

# Extensions we can extract text from.
_EXTRACTABLE_EXTENSIONS: frozenset[str] = frozenset({
    "xlsx", "xls", "docx", "csv", "txt", "html", "htm", "eml",
})


class TextPart(NamedTuple):
    filename: str
    text: str


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def should_skip(filename: str) -> bool:
    """Return True for junk/system files — skip upload filtering too."""
    basename = filename.rsplit("/", 1)[-1]
    if basename in _SKIP_BASENAMES:
        return True
    ext = basename.rsplit(".", 1)[-1].lower() if "." in basename else ""
    return ext in _SKIP_EXTENSIONS


def extractable_extension(filename: str) -> str | None:
    """Return lowercase extension if we support text extraction, else None."""
    if "." not in filename:
        return None
    ext = filename.rsplit(".", 1)[-1].lower()
    return ext if ext in _EXTRACTABLE_EXTENSIONS else None


# ---------------------------------------------------------------------------
# Per-format extractors (all synchronous; called via asyncio.to_thread)
# ---------------------------------------------------------------------------


def _is_html_content(content: bytes) -> bool:
    """Return True if the bytes look like HTML (some banks save HTML as .xls/.xlsx)."""
    sniff = content[:512].lstrip()
    return sniff.startswith(b"<") or b"<html" in sniff.lower() or b"<HTML" in sniff


def _xlsx_to_text(content: bytes, filename: str) -> str:
    import openpyxl  # type: ignore[import]

    if _is_html_content(content):
        return _html_to_text(content, filename)

    wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True, read_only=True)
    parts = [f"=== Excel document: {filename} ==="]
    per_sheet_budget = MAX_TEXT_CHARS // max(1, len(wb.sheetnames))
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_lines: list[str] = []
        chars = 0
        for row in ws.iter_rows(values_only=True):
            if any(c is not None for c in row):
                line = "\t".join("" if c is None else str(c) for c in row)
                sheet_lines.append(line)
                chars += len(line)
                if chars >= per_sheet_budget:
                    sheet_lines.append("... [sheet truncated] ...")
                    break
        if sheet_lines:
            parts.append(f"\n--- Sheet: {sheet_name} ---")
            parts.extend(sheet_lines)
    wb.close()
    return "\n".join(parts)[:MAX_TEXT_CHARS]


def _xls_to_text(content: bytes, filename: str) -> str:
    import xlrd  # type: ignore[import]

    if _is_html_content(content):
        return _html_to_text(content, filename)

    wb = xlrd.open_workbook(file_contents=content)
    parts = [f"=== Excel document (legacy .xls): {filename} ==="]
    per_sheet_budget = MAX_TEXT_CHARS // max(1, wb.nsheets)
    for sheet_name in wb.sheet_names():
        ws = wb.sheet_by_name(sheet_name)
        parts.append(f"\n--- Sheet: {sheet_name} ---")
        chars = 0
        for row_idx in range(ws.nrows):
            row = ws.row_values(row_idx)
            if any(c != "" for c in row):
                line = "\t".join(str(c) for c in row)
                parts.append(line)
                chars += len(line)
                if chars >= per_sheet_budget:
                    parts.append("... [sheet truncated] ...")
                    break
    return "\n".join(parts)[:MAX_TEXT_CHARS]


def _docx_to_text(content: bytes, filename: str) -> str:
    import docx as python_docx  # type: ignore[import]

    doc = python_docx.Document(io.BytesIO(content))
    parts = [f"=== Word document: {filename} ===\n"]
    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            parts.append(stripped)
    for table in doc.tables:
        parts.append("")
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)[:MAX_TEXT_CHARS]


def _csv_to_text(content: bytes, filename: str) -> str:
    text = content.decode("utf-8", errors="replace")
    return f"=== CSV document: {filename} ===\n{text[:MAX_TEXT_CHARS]}"


def _txt_to_text(content: bytes, filename: str) -> str:
    text = content.decode("utf-8", errors="replace")
    return f"=== Text document: {filename} ===\n{text[:MAX_TEXT_CHARS]}"


def _html_to_text(content: bytes, filename: str) -> str:
    from html.parser import HTMLParser

    class _TextExtractor(HTMLParser):
        def __init__(self) -> None:
            super().__init__()
            self._parts: list[str] = []
            self._skip = False

        def handle_starttag(self, tag: str, attrs: list) -> None:
            if tag in ("script", "style"):
                self._skip = True

        def handle_endtag(self, tag: str) -> None:
            if tag in ("script", "style"):
                self._skip = False

        def handle_data(self, data: str) -> None:
            if not self._skip and data.strip():
                self._parts.append(data.strip())

    extractor = _TextExtractor()
    extractor.feed(content.decode("utf-8", errors="replace"))
    text = "\n".join(extractor._parts)
    return f"=== HTML document: {filename} ===\n{text[:MAX_TEXT_CHARS]}"


def _eml_to_text(content: bytes, filename: str) -> str:
    import email as email_lib

    msg = email_lib.message_from_bytes(content)
    subject = msg.get("Subject", "")
    sender = msg.get("From", "")
    date = msg.get("Date", "")

    body_parts: list[str] = []
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == "text/plain":
                payload = part.get_payload(decode=True)
                if payload:
                    body_parts.append(payload.decode("utf-8", errors="replace"))
    else:
        payload = msg.get_payload(decode=True)
        if payload:
            body_parts.append(payload.decode("utf-8", errors="replace"))

    text = (
        f"=== Email: {filename} ===\n"
        f"From: {sender}\nDate: {date}\nSubject: {subject}\n\n"
        + "\n".join(body_parts)
    )
    return text[:MAX_TEXT_CHARS]


def _extract_bytes(content: bytes, filename: str, ext: str) -> str | None:
    """Dispatch to the correct extractor by extension."""
    try:
        if ext == "xlsx":
            return _xlsx_to_text(content, filename)
        if ext == "xls":
            return _xls_to_text(content, filename)
        if ext == "docx":
            return _docx_to_text(content, filename)
        if ext == "csv":
            return _csv_to_text(content, filename)
        if ext == "txt":
            return _txt_to_text(content, filename)
        if ext in ("html", "htm"):
            return _html_to_text(content, filename)
        if ext == "eml":
            return _eml_to_text(content, filename)
    except Exception:
        logger.warning("Text extraction failed for %s", filename, exc_info=True)
    return None


# ---------------------------------------------------------------------------
# GCS download + extraction (blocking — run via asyncio.to_thread)
# ---------------------------------------------------------------------------


def _parse_gcs_uri(gcs_uri: str) -> tuple[str, str]:
    """Parse ``gs://bucket/object/path`` → ``(bucket, object_path)``."""
    if not gcs_uri.startswith("gs://"):
        raise ValueError(f"Not a GCS URI: {gcs_uri!r}")
    rest = gcs_uri[5:]
    bucket, _, obj = rest.partition("/")
    return bucket, obj


def _download_and_extract(
    gcs_uri: str,
    filename: str,
    ext: str,
    file_size_bytes: int | None,
) -> TextPart | None:
    """Download a file from GCS and extract its text content.

    Designed to be called via ``asyncio.to_thread`` so it doesn't block the
    event loop.  Returns ``None`` when the file is too large, can't be
    downloaded, or text extraction fails.
    """
    if file_size_bytes and file_size_bytes > MAX_EXTRACT_BYTES:
        logger.warning(
            "text_extractor: skipping %s — size %d bytes exceeds limit %d",
            filename,
            file_size_bytes,
            MAX_EXTRACT_BYTES,
        )
        return None

    try:
        from google.cloud import storage  # type: ignore[import]

        bucket_name, obj_path = _parse_gcs_uri(gcs_uri)
        client = storage.Client()
        bucket = client.bucket(bucket_name)
        blob = bucket.blob(obj_path)

        if file_size_bytes is None:
            blob.reload()
            actual_size = blob.size or 0
            if actual_size > MAX_EXTRACT_BYTES:
                logger.warning(
                    "text_extractor: skipping %s — GCS size %d exceeds limit",
                    filename,
                    actual_size,
                )
                return None

        content = blob.download_as_bytes()
    except Exception:
        logger.warning(
            "text_extractor: GCS download failed for %s (%s)",
            filename,
            gcs_uri,
            exc_info=True,
        )
        return None

    text = _extract_bytes(content, filename, ext)
    if text:
        logger.info(
            "text_extractor: extracted %d chars from %s", len(text), filename
        )
        return TextPart(filename=filename, text=text)
    return None


# ---------------------------------------------------------------------------
# Public async API
# ---------------------------------------------------------------------------


async def extract_text_parts(files: list) -> list[TextPart]:
    """Download and extract text from non-Gemini files concurrently.

    ``files`` is a list of SQLAlchemy ``File`` ORM objects (or any object with
    ``original_name``, ``gcs_uri``, and ``file_size_bytes`` attributes).

    Returns a list of ``TextPart`` namedtuples (filename, text) in arbitrary
    order.  Files that are too large, fail to download, or whose format we
    can't handle are silently skipped (a WARNING is logged).
    """
    tasks: list[tuple] = []
    for f in files:
        fname = f.original_name
        if should_skip(fname):
            logger.debug("text_extractor: skipping system file %s", fname)
            continue
        ext = extractable_extension(fname)
        if ext is None:
            logger.debug("text_extractor: no extractor for %s", fname)
            continue
        tasks.append((f, ext))

    if not tasks:
        return []

    coros = [
        asyncio.to_thread(
            _download_and_extract,
            f.gcs_uri,
            f.original_name,
            ext,
            getattr(f, "file_size_bytes", None),
        )
        for f, ext in tasks
    ]
    outcomes = await asyncio.gather(*coros, return_exceptions=True)

    results: list[TextPart] = []
    for outcome in outcomes:
        if isinstance(outcome, TextPart):
            results.append(outcome)
        elif isinstance(outcome, Exception):
            logger.warning("text_extractor: extraction task error: %s", outcome)
        # None outcomes mean skipped/failed — already logged inside the sync fn

    logger.info(
        "text_extractor: %d/%d files yielded text", len(results), len(tasks)
    )
    return results
