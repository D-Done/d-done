"""Word document export service for DD reports.

Generates a Hebrew RTL .docx file from a RealEstateFinanceDDReport or DDReport.
Uses python-docx with manual XML tweaks for right-to-left paragraph support.
In-app comments are injected as proper Word margin comments via ZIP post-processing.
"""

from __future__ import annotations

import io
import zipfile
from datetime import date, datetime, timezone
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

from app.agents.schemas import (
    DDReport,
    RealEstateFinanceDDReport,
    Finding,
    TenantRow,
)

# ── XML namespaces ────────────────────────────────────────────────────────────

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_REL = "http://schemas.openxmlformats.org/package/2006/relationships"
_CT = "http://schemas.openxmlformats.org/package/2006/content-types"

# Maps section_key → heading text that appears verbatim in the DOCX
SECTION_HEADING_TEXTS: dict[str, str] = {
    "executive_summary": "1. סיכום מנהלים",
    "compound_details": "3. פרטי המתחם",
    "tenant_table": "4. טבלת דיירים",
    "developer_signature": "5. חתימת היזם",
    "legal_representation": "6. באי כוח",
    "financing_body": "7. גוף המימון",
    "zero_report": '8. דו"ח אפס',
    "corporate_governance": "10. שרשרת בעלות (UBO)",
    "findings": "12. ממצאים",
    # Standard report
    "findings_std": "3. ממצאים",
}

# ---------------------------------------------------------------------------
# RTL helpers
# ---------------------------------------------------------------------------

RISK_LABEL_HE = {"high": "גבוה", "medium": "בינוני", "low": "נמוך"}
SEVERITY_LABEL_HE = {"critical": "קריטי", "warning": "אזהרה", "info": "מידע"}
BOOL_HE = {True: "כן", False: "לא", None: "—"}

RISK_LABEL_EN = {"high": "High", "medium": "Medium", "low": "Low"}
SEVERITY_LABEL_EN = {"critical": "Critical", "warning": "Warning", "info": "Info"}
BOOL_EN = {True: "Yes", False: "No", None: "—"}

# Section headings and field labels for English
_EN_LABELS: dict[str, str] = {
    # Cover
    "דוח בדיקת נאותות": "Due Diligence Report",
    "תאריך הפקה": "Date",
    "לקוח": "Client",
    "סטטוס": "Status",
    "מסמכים שנותחו": "Documents Analyzed",
    # Sections
    "1. סיכום מנהלים": "1. Executive Summary",
    "3. פרטי המתחם": "3. Compound Details",
    "4. טבלת דיירים": "4. Tenant Table",
    "5. חתימת היזם": "5. Developer Signature",
    "6. באי כוח": "6. Legal Representatives",
    "7. גוף המימון": "7. Financing Body",
    '8. דו"ח אפס': "8. Zero Report",
    "9. ניתוח פיננסי": "9. Financial Analysis",
    "10. שרשרת בעלות (UBO)": "10. Ownership Chain (UBO)",
    "קשרי בעלות": "Ownership Relations",
    "11. דגלים אדומים": "11. Red Flags",
    "12. ממצאים": "12. Findings",
    "2. סיכום מנהלים": "2. Executive Summary",
    "3. ממצאים": "3. Findings",
    # Finance fields
    "רמת סיכון": "Risk Level",
    "סיכום": "Summary",
    "המלצה": "Recommendation",
    "סיכונים עיקריים": "Key Risks",
    "כתובת": "Address",
    "גוש": "Block",
    "חלקה": "Parcel",
    "מצב לפני הריסה": "Pre-demolition State",
    "מצב לאחר בנייה": "Post-construction State",
    "פערים": "Discrepancies",
    "בניינים": "buildings",
    "דירות": "apartments",
    "אחוז חתימות": "Signing Rate",
    "תת-חלקה": "Sub-parcel",
    "שם בעלים": "Owner Name",
    "חתם": "Signed",
    "תאריך חתימה": "Date Signed",
    "הערת אזהרה": "Caveat",
    "משכנתא": "Mortgage",
    "הערות": "Notes",
    "תאריך חתימת יזם": "Developer Signing Date",
    "שם מורשה חתימה": "Authorized Signatory Name",
    "ת.ז. מורשה חתימה": "Authorized Signatory ID",
    "אישור פרוטוקול חתימה": "Signing Protocol Approval",
    "מאושר": "Approved",
    "אי-התאמה": "Mismatch",
    "בא כוח היזם": "Developer's Attorney",
    "בא כוח הבעלים": "Owners' Attorney",
    "הגדרת מממן בהסכם": "Lender Definition in Agreement",
    "מממן בפועל": "Actual Lender",
    "עמידה בתנאים": "Compliance Note",
    "הלוואת מזנין": "Mezzanine Loan",
    "פרטי מזנין": "Mezzanine Details",
    "נמען הדוח": "Report Addressee",
    "רווח למחזור": "Profit on Turnover",
    "רווח לעלות": "Profit on Cost",
    "הצמדה למדד": "Indexation",
    "מגבלות בנייה": "Construction Restrictions",
    "התאמת הגדרת מממן": "Lender Definition Match",
    "פרטי אי-התאמה": "Discrepancy Details",
    "אישור הון עצמי": "Equity Confirmation",
    "תואם": "Compliant",
    "בעלים": "Owner",
    "חברה": "Company",
    "אחוז החזקה": "Stake %",
    # Standard report
    "ממצאים": "Findings",
    "מידע כללי": "General Information",
    "שדה": "Field",
    "ערך": "Value",
}


def _lbl(text: str, language: str) -> str:
    """Return English label if language=='en', else return Hebrew as-is."""
    if language == "en":
        return _EN_LABELS.get(text, text)
    return text



def _set_rtl_paragraph(para: Any) -> None:
    """Add RTL + right-alignment to an existing paragraph."""
    pPr = para._p.get_or_add_pPr()
    _ensure_bidi(pPr)
    _ensure_jc_right(pPr)


def _set_rtl_run(run: Any) -> None:
    """Add <w:rtl/> to a run's rPr so the font renders RTL."""
    rPr = run._r.get_or_add_rPr()
    rtl = rPr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        rPr.append(rtl)


def _add_rtl_paragraph(doc: Document, text: str, style: str | None = None) -> Any:
    """Add a new RTL paragraph with the given text and optional style."""
    if style:
        para = doc.add_paragraph(style=style)
    else:
        para = doc.add_paragraph()
    _set_rtl_paragraph(para)
    run = para.add_run(text)
    _set_rtl_run(run)
    return para


def _add_heading(doc: Document, text: str, level: int = 1) -> Any:
    """Add an RTL heading."""
    para = doc.add_heading(level=level)
    para.clear()
    _set_rtl_paragraph(para)
    run = para.add_run(text)
    _set_rtl_run(run)
    return para


def _add_bold_label(doc: Document, label: str, value: str | None) -> None:
    """Add a paragraph like: **label:** value."""
    if not value:
        return
    para = doc.add_paragraph()
    _set_rtl_paragraph(para)
    bold_run = para.add_run(f"{label}: ")
    bold_run.bold = True
    _set_rtl_run(bold_run)
    val_run = para.add_run(value)
    _set_rtl_run(val_run)


def _set_table_rtl(table: Any) -> None:
    """Set RTL on a whole table's tblPr."""
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    bidi = tblPr.find(qn("w:bidiVisual"))
    if bidi is None:
        bidi = OxmlElement("w:bidiVisual")
        tblPr.append(bidi)


def _cell_rtl(cell: Any, text: str, bold: bool = False) -> None:
    """Set cell text with RTL alignment."""
    cell.text = ""
    para = cell.paragraphs[0]
    _set_rtl_paragraph(para)
    run = para.add_run(text)
    _set_rtl_run(run)
    if bold:
        run.bold = True


# ---------------------------------------------------------------------------
# Finance report builder
# ---------------------------------------------------------------------------


def _build_finance_report(doc: Document, report: RealEstateFinanceDDReport, project_title: str, language: str = "he") -> None:
    bool_map = BOOL_EN if language == "en" else BOOL_HE
    risk_map = RISK_LABEL_EN if language == "en" else RISK_LABEL_HE

    # ---- Cover ----
    title_para = doc.add_heading(level=0)
    title_para.clear()
    _set_rtl_paragraph(title_para)
    run = title_para.add_run(f"{_lbl('דוח בדיקת נאותות', language)} — {project_title}")
    _set_rtl_run(run)

    date_str = date.today().strftime("%d/%m/%Y")
    _add_rtl_paragraph(doc, f"{_lbl('תאריך הפקה', language)}: {date_str}")

    if report.project_header:
        h = report.project_header
        if h.client_name:
            _add_rtl_paragraph(doc, f"{_lbl('לקוח', language)}: {h.client_name}")
        if h.status:
            _add_rtl_paragraph(doc, f"{_lbl('סטטוס', language)}: {h.status}")
        if h.doc_count is not None:
            _add_rtl_paragraph(doc, f"{_lbl('מסמכים שנותחו', language)}: {h.doc_count}")

    doc.add_page_break()

    # ---- 1. Executive Summary ----
    _add_heading(doc, _lbl("1. סיכום מנהלים", language), level=1)
    es = report.executive_summary
    risk_label = risk_map.get(es.risk_level, es.risk_level)
    _add_bold_label(doc, _lbl("רמת סיכון", language), risk_label)
    _add_bold_label(doc, _lbl("סיכום", language), es.summary)
    if rec := getattr(es, "recommendation", None):
        _add_bold_label(doc, _lbl("המלצה", language), rec)

    if getattr(es, "key_risks", None):
        para = doc.add_paragraph()
        _set_rtl_paragraph(para)
        r = para.add_run(f"{_lbl('סיכונים עיקריים', language)}:")
        r.bold = True
        _set_rtl_run(r)
        for risk in getattr(es, "key_risks", []):
            p = doc.add_paragraph(style="List Bullet")
            _set_rtl_paragraph(p)
            run = p.add_run(risk)
            _set_rtl_run(run)

    # ---- 2. Compound Details ----
    if report.compound_details:
        _add_heading(doc, _lbl("3. פרטי המתחם", language), level=1)
        cd = report.compound_details
        _add_bold_label(doc, _lbl("כתובת", language), cd.address)
        _add_bold_label(doc, _lbl("גוש", language), cd.gush)
        _add_bold_label(doc, _lbl("חלקה", language), cd.helka)
        bldg = _lbl("בניינים", language)
        apts = _lbl("דירות", language)
        if cd.incoming_state:
            s = cd.incoming_state
            _add_bold_label(doc, _lbl("מצב לפני הריסה", language),
                            f"{s.building_count or '—'} {bldg}, {s.apartment_count or '—'} {apts}")
        if cd.outgoing_state:
            s = cd.outgoing_state
            _add_bold_label(doc, _lbl("מצב לאחר בנייה", language),
                            f"{s.building_count or '—'} {bldg}, {s.apartment_count or '—'} {apts}")
        _add_bold_label(doc, _lbl("פערים", language), cd.discrepancy_note)

    # ---- 4. Tenant Table ----
    if report.tenant_table:
        _add_heading(doc, _lbl("4. טבלת דיירים", language), level=1)

        signing_pct = report.signing_percentage or 0
        if signing_pct <= 1:
            signing_pct = round(signing_pct * 100)
        else:
            signing_pct = round(signing_pct)
        _add_rtl_paragraph(doc, f"{_lbl('אחוז חתימות', language)}: {signing_pct}%")

        cols = [
            _lbl("תת-חלקה", language), _lbl("שם בעלים", language), _lbl("חתם", language),
            _lbl("תאריך חתימה", language), _lbl("הערת אזהרה", language),
            _lbl("משכנתא", language), _lbl("הערות", language),
        ]
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"
        _set_table_rtl(table)
        hdr = table.rows[0].cells
        for i, col in enumerate(cols):
            _cell_rtl(hdr[i], col, bold=True)

        for row_data in report.tenant_table:
            row = table.add_row().cells
            _cell_rtl(row[0], row_data.sub_parcel or row_data.helka or "")
            _cell_rtl(row[1], row_data.owner_name or "")
            _cell_rtl(row[2], bool_map[row_data.is_signed])
            _cell_rtl(row[3], row_data.date_signed or "")
            _cell_rtl(row[4], bool_map[row_data.is_warning_note_registered])
            _cell_rtl(row[5], bool_map[row_data.is_mortgage_registered])
            _cell_rtl(row[6], row_data.notes or "")
        doc.add_paragraph()

    # ---- 5. Developer Signature ----
    if report.developer_signature:
        _add_heading(doc, _lbl("5. חתימת היזם", language), level=1)
        ds = report.developer_signature
        _add_bold_label(doc, _lbl("תאריך חתימת יזם", language), ds.developer_signed_date)
        _add_bold_label(doc, _lbl("שם מורשה חתימה", language), ds.authorized_signatory_name)
        _add_bold_label(doc, _lbl("ת.ז. מורשה חתימה", language), ds.authorized_signatory_id)
        if ds.signing_protocol_authorized is not None:
            _add_bold_label(doc, _lbl("אישור פרוטוקול חתימה", language),
                            _lbl("מאושר", language) if ds.signing_protocol_authorized else _lbl("אי-התאמה", language))

    # ---- 6. Power of Attorney ----
    if report.power_of_attorney:
        _add_heading(doc, _lbl("6. באי כוח", language), level=1)
        poa = report.power_of_attorney
        _add_bold_label(doc, _lbl("בא כוח היזם", language), poa.developer_attorney)
        _add_bold_label(doc, _lbl("בא כוח הבעלים", language), poa.owners_attorney)

    # ---- 7. Financing Body ----
    if report.financing:
        _add_heading(doc, _lbl("7. גוף המימון", language), level=1)
        fin = report.financing
        _add_bold_label(doc, _lbl("הגדרת מממן בהסכם", language), fin.lender_definition_clause)
        _add_bold_label(doc, _lbl("מממן בפועל", language), fin.actual_lender)
        _add_bold_label(doc, _lbl("עמידה בתנאים", language), fin.lender_compliance_note)
        if fin.mezzanine_loan_exists is not None:
            _add_bold_label(doc, _lbl("הלוואת מזנין", language), bool_map[fin.mezzanine_loan_exists])
        _add_bold_label(doc, _lbl("פרטי מזנין", language), fin.mezzanine_loan_details)

    # ---- 8. Zero Report Metrics ----
    if report.zero_report_metrics:
        _add_heading(doc, _lbl('8. דו"ח אפס', language), level=1)
        zr = report.zero_report_metrics
        _add_bold_label(doc, _lbl("נמען הדוח", language), zr.addressee)
        if zr.profit_on_turnover is not None:
            _add_bold_label(doc, _lbl("רווח למחזור", language), f"{zr.profit_on_turnover:.1%}")
        if zr.profit_on_cost is not None:
            _add_bold_label(doc, _lbl("רווח לעלות", language), f"{zr.profit_on_cost:.1%}")
        _add_bold_label(doc, _lbl("הצמדה למדד", language), zr.indexation_details)
        if zr.construction_restrictions:
            para = doc.add_paragraph()
            _set_rtl_paragraph(para)
            r = para.add_run(f"{_lbl('מגבלות בנייה', language)}:")
            r.bold = True
            _set_rtl_run(r)
            for restriction in zr.construction_restrictions:
                p = doc.add_paragraph(style="List Bullet")
                _set_rtl_paragraph(p)
                run = p.add_run(restriction)
                _set_rtl_run(run)

    # ---- 9. Finance Analysis ----
    if getattr(report, "finance_analysis", None):
        _add_heading(doc, _lbl("9. ניתוח פיננסי", language), level=1)
        fa = report.finance_analysis
        if fa.lender_definition_match is not None:
            _add_bold_label(doc, _lbl("התאמת הגדרת מממן", language),
                            _lbl("תואם", language) if fa.lender_definition_match else _lbl("אי-התאמה", language))
        _add_bold_label(doc, _lbl("פרטי אי-התאמה", language), fa.discrepancy_note)
        if fa.equity_confirmed is not None:
            _add_bold_label(doc, _lbl("אישור הון עצמי", language), bool_map[fa.equity_confirmed])

    # ---- 10. Corporate Governance / UBO ----
    if report.developer_ubo_chain:
        _add_heading(doc, _lbl("10. שרשרת בעלות (UBO)", language), level=1)
        for item in report.developer_ubo_chain:
            p = doc.add_paragraph(style="List Bullet")
            _set_rtl_paragraph(p)
            run = p.add_run(item)
            _set_rtl_run(run)

    if report.developer_ubo_graph and report.developer_ubo_graph.edges:
        _add_heading(doc, _lbl("קשרי בעלות", language), level=2)
        graph = report.developer_ubo_graph
        node_map = {n.id: n.name for n in graph.nodes}
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        _set_table_rtl(table)
        hdr = table.rows[0].cells
        _cell_rtl(hdr[0], _lbl("בעלים", language), bold=True)
        _cell_rtl(hdr[1], _lbl("חברה", language), bold=True)
        _cell_rtl(hdr[2], _lbl("אחוז החזקה", language), bold=True)
        for edge in graph.edges:
            row = table.add_row().cells
            _cell_rtl(row[0], node_map.get(edge.from_id, edge.from_id))
            _cell_rtl(row[1], node_map.get(edge.to_id, edge.to_id))
            _cell_rtl(row[2], edge.share_pct or "—")
        doc.add_paragraph()

    # ---- 11. High Risk Flags ----
    if report.high_risk_flags:
        _add_heading(doc, _lbl("11. דגלים אדומים", language), level=1)
        for flag in report.high_risk_flags:
            p = doc.add_paragraph(style="List Bullet")
            _set_rtl_paragraph(p)
            run = p.add_run(flag)
            _set_rtl_run(run)

    # ---- 12. Findings ----
    if report.findings:
        _add_heading(doc, _lbl("12. ממצאים", language), level=1)
        _build_findings(doc, report.findings, language=language)


# ---------------------------------------------------------------------------
# Standard DDReport builder
# ---------------------------------------------------------------------------


def _build_standard_report(doc: Document, report: DDReport, project_title: str, language: str = "he") -> None:
    risk_map = RISK_LABEL_EN if language == "en" else RISK_LABEL_HE

    title_para = doc.add_heading(level=0)
    title_para.clear()
    _set_rtl_paragraph(title_para)
    run = title_para.add_run(f"{_lbl('דוח בדיקת נאותות', language)} — {project_title}")
    _set_rtl_run(run)

    date_str = date.today().strftime("%d/%m/%Y")
    _add_rtl_paragraph(doc, f"{_lbl('תאריך הפקה', language)}: {date_str}")
    doc.add_page_break()

    # Executive Summary
    _add_heading(doc, _lbl("1. סיכום מנהלים", language), level=1)
    es = report.executive_summary
    _add_bold_label(doc, _lbl("רמת סיכון", language), risk_map.get(es.risk_level, es.risk_level))
    _add_bold_label(doc, _lbl("סיכום", language), es.summary)
    if rec := getattr(es, "recommendation", None):
        _add_bold_label(doc, _lbl("המלצה", language), rec)
    if getattr(es, "key_risks", None):
        para = doc.add_paragraph()
        _set_rtl_paragraph(para)
        r = para.add_run(f"{_lbl('סיכונים עיקריים', language)}:")
        r.bold = True
        _set_rtl_run(r)
        for risk in getattr(es, "key_risks", []):
            p = doc.add_paragraph(style="List Bullet")
            _set_rtl_paragraph(p)
            run = p.add_run(risk)
            _set_rtl_run(run)

    # Findings
    if report.findings:
        _add_heading(doc, _lbl("3. ממצאים", language), level=1)
        _build_findings(doc, report.findings, language=language)

    # Documents Analyzed
    if report.documents_analyzed:
        pages_word = "pages" if language == "en" else "עמודים"
        _add_heading(doc, _lbl("מסמכים שנותחו", language), level=1)
        for doc_item in report.documents_analyzed:
            _add_bold_label(doc, doc_item.name, f"{doc_item.page_count} {pages_word}")


# ---------------------------------------------------------------------------
# Shared findings renderer
# ---------------------------------------------------------------------------


def _build_findings(doc: Document, findings: list[Finding], language: str = "he") -> None:
    sev_map = SEVERITY_LABEL_EN if language == "en" else SEVERITY_LABEL_HE
    sources_lbl = "Sources: " if language == "en" else "מקורות: "
    page_abbr = "p." if language == "en" else "עמ'"

    for f in findings:
        severity_label = sev_map.get(f.severity, f.severity)
        para = doc.add_heading(level=3)
        para.clear()
        _set_rtl_paragraph(para)
        run = para.add_run(f"[{severity_label}] {f.title}")
        _set_rtl_run(run)

        desc_para = doc.add_paragraph()
        _set_rtl_paragraph(desc_para)
        run = desc_para.add_run(f.description)
        _set_rtl_run(run)

        if f.sources:
            src_para = doc.add_paragraph()
            _set_rtl_paragraph(src_para)
            label_run = src_para.add_run(sources_lbl)
            label_run.bold = True
            _set_rtl_run(label_run)
            for i, src in enumerate(f.sources):
                if i > 0:
                    sep_run = src_para.add_run(" | ")
                    _set_rtl_run(sep_run)
                src_run = src_para.add_run(
                    f"{src.source_document_name} {page_abbr} {src.page_number}"
                )
                _set_rtl_run(src_run)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


# OOXML schema: elements that must appear BEFORE <w:bidi> inside <w:pPr>
_PPR_BEFORE_BIDI = {
    qn("w:pStyle"), qn("w:keepNext"), qn("w:keepLines"),
    qn("w:pageBreakBefore"), qn("w:framePr"), qn("w:suppressLineNumbers"),
    qn("w:pBdr"), qn("w:shd"), qn("w:tabs"), qn("w:suppressAutoHyphens"),
    qn("w:kinsoku"), qn("w:wordWrap"), qn("w:overflowPunct"),
    qn("w:topLinePunct"), qn("w:autoSpaceDE"), qn("w:autoSpaceDN"),
}


def _ensure_bidi(pPr: Any) -> None:
    """Insert <w:bidi/> at the correct position in a <w:pPr> element."""
    if pPr.find(qn("w:bidi")) is not None:
        return
    bidi = OxmlElement("w:bidi")
    # Insert after the last element that must precede <w:bidi>
    insert_idx = 0
    for i, child in enumerate(pPr):
        if child.tag in _PPR_BEFORE_BIDI:
            insert_idx = i + 1
    pPr.insert(insert_idx, bidi)


def _ensure_jc_right(pPr: Any) -> None:
    """Set <w:jc w:val="right"/> in a <w:pPr>, replacing any existing value."""
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), "right")


def _configure_document_rtl(doc: Document) -> None:
    """Set RTL at the document-defaults and styles level so every element inherits it.

    <w:docDefaults> lives in styles.xml (doc.styles.element), NOT settings.xml.
    """
    styles_element = doc.styles.element

    # 1. Document defaults (<w:docDefaults> is a child of <w:styles>)
    docDefaults = styles_element.find(qn("w:docDefaults"))
    if docDefaults is not None:
        # pPrDefault → <w:bidi/> + right-align
        pPrDefault = docDefaults.find(qn("w:pPrDefault"))
        if pPrDefault is None:
            pPrDefault = OxmlElement("w:pPrDefault")
            docDefaults.insert(0, pPrDefault)
        pPr = pPrDefault.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            pPrDefault.append(pPr)
        _ensure_bidi(pPr)
        _ensure_jc_right(pPr)

        # rPrDefault → <w:rtl/>
        rPrDefault = docDefaults.find(qn("w:rPrDefault"))
        if rPrDefault is None:
            rPrDefault = OxmlElement("w:rPrDefault")
            docDefaults.append(rPrDefault)
        rPr = rPrDefault.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            rPrDefault.append(rPr)
        if rPr.find(qn("w:rtl")) is None:
            rPr.append(OxmlElement("w:rtl"))

    # 2. Patch every style so they inherit/enforce RTL
    for style in styles_element.findall(qn("w:style")):
        pPr = style.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            style.append(pPr)
        _ensure_bidi(pPr)
        _ensure_jc_right(pPr)

        rPr = style.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            style.append(rPr)
        if rPr.find(qn("w:rtl")) is None:
            rPr.append(OxmlElement("w:rtl"))


def generate_word_report(
    report: DDReport | RealEstateFinanceDDReport,
    project_title: str,
    comments_by_section: dict[str, list[dict]] | None = None,
    language: str = "he",
) -> bytes:
    """Generate a .docx file from a DD report and return its bytes.

    comments_by_section: {section_key: [{content, author_name, author_email, created_at}, ...]}
    If provided, comments are injected as proper Word margin comments.
    """
    doc = Document()
    _configure_document_rtl(doc)

    if isinstance(report, RealEstateFinanceDDReport):
        _build_finance_report(doc, report, project_title, language=language)
    else:
        _build_standard_report(doc, report, project_title, language=language)

    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    if comments_by_section:
        docx_bytes = _inject_word_comments(docx_bytes, comments_by_section)

    return docx_bytes


# ── Word comment injection (ZIP post-processing) ──────────────────────────────

def _inject_word_comments(
    docx_bytes: bytes,
    comments_by_section: dict[str, list[dict]],
) -> bytes:
    """
    Post-process the DOCX ZIP to embed proper Word comment XML.
    Each comment appears as a margin annotation on the matching section heading.
    """
    # Build text→comments map using known heading labels
    label_to_comments: dict[str, list[dict]] = {}
    for section_key, cmts in comments_by_section.items():
        label = SECTION_HEADING_TEXTS.get(section_key)
        if label and cmts:
            label_to_comments[label] = cmts

    if not label_to_comments:
        return docx_bytes

    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zin:
        all_files = {info.filename: zin.read(info.filename) for info in zin.infolist()}

    doc_tree = etree.fromstring(all_files["word/document.xml"])

    comments_list: list[tuple[int, str, str, str]] = []
    comment_id = 0

    for para in doc_tree.iter(f"{{{_W}}}p"):
        para_text = "".join(t.text or "" for t in para.iter(f"{{{_W}}}t"))
        if para_text not in label_to_comments:
            continue
        for cmt in label_to_comments[para_text]:
            cid = comment_id
            comment_id += 1

            cs = etree.Element(f"{{{_W}}}commentRangeStart")
            cs.set(f"{{{_W}}}id", str(cid))
            para.insert(0, cs)

            ce = etree.Element(f"{{{_W}}}commentRangeEnd")
            ce.set(f"{{{_W}}}id", str(cid))
            para.append(ce)

            ref_run = etree.Element(f"{{{_W}}}r")
            rpr = etree.SubElement(ref_run, f"{{{_W}}}rPr")
            rs = etree.SubElement(rpr, f"{{{_W}}}rStyle")
            rs.set(f"{{{_W}}}val", "CommentReference")
            comment_ref = etree.SubElement(ref_run, f"{{{_W}}}commentReference")
            comment_ref.set(f"{{{_W}}}id", str(cid))
            para.append(ref_run)

            author = cmt.get("author_name") or cmt.get("author_email") or "Reviewer"
            ca = cmt.get("created_at")
            if isinstance(ca, datetime):
                date_str = ca.strftime("%Y-%m-%dT%H:%M:%SZ")
            elif ca:
                date_str = str(ca)
            else:
                date_str = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

            comments_list.append((cid, author, date_str, cmt["content"]))

    # Build comments.xml
    croot = etree.Element(f"{{{_W}}}comments", nsmap={"w": _W, "r": _R})
    for cid, author, date_str, text in comments_list:
        ce = etree.SubElement(croot, f"{{{_W}}}comment")
        ce.set(f"{{{_W}}}id", str(cid))
        ce.set(f"{{{_W}}}author", author)
        ce.set(f"{{{_W}}}date", date_str)
        cp = etree.SubElement(ce, f"{{{_W}}}p")
        cr = etree.SubElement(cp, f"{{{_W}}}r")
        ct_elem = etree.SubElement(cr, f"{{{_W}}}t")
        ct_elem.text = text

    all_files["word/document.xml"] = etree.tostring(
        doc_tree, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    all_files["word/comments.xml"] = etree.tostring(
        croot, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    # Add comments relationship
    rels_raw = all_files.get("word/_rels/document.xml.rels", b"<Relationships/>")
    rels_tree = etree.fromstring(rels_raw)
    rel = etree.SubElement(rels_tree, f"{{{_REL}}}Relationship")
    rel.set("Id", "rIdComments")
    rel.set("Type", f"{_R}/comments")
    rel.set("Target", "comments.xml")
    all_files["word/_rels/document.xml.rels"] = etree.tostring(rels_tree)

    # Add content type override
    ct_tree = etree.fromstring(all_files["[Content_Types].xml"])
    ov = etree.SubElement(ct_tree, f"{{{_CT}}}Override")
    ov.set("PartName", "/word/comments.xml")
    ov.set("ContentType",
           "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml")
    all_files["[Content_Types].xml"] = etree.tostring(ct_tree)

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for filename, data in all_files.items():
            zout.writestr(filename, data)
    return out.getvalue()
